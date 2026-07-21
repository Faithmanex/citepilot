import html
import io
import logging
from typing import Dict

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

logger = logging.getLogger(__name__)


def generate_pdf_report(analysis_data: Dict) -> bytes:
    """
    Generates a PDF Diagnostic Report with ReportLab or minimal valid PDF fallback.
    Wraps dynamic text in html.escape() to prevent ReportLab Expat XML parser errors.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=20,
            leading=24,
            textColor=colors.HexColor('#1E3A8A')
        )

        h2_style = ParagraphStyle(
            'DocH2',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=14,
            leading=18,
            textColor=colors.HexColor('#0F172A'),
            spaceBefore=12,
            spaceAfter=6
        )

        body_style = ParagraphStyle(
            'DocBody',
            parent=styles['BodyText'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#334155')
        )

        story = []
        story.append(Paragraph("CitePilot Diagnostic Citation Audit Report", title_style))
        story.append(Spacer(1, 12))

        citations = analysis_data.get("citations", [])
        refs = analysis_data.get("references", [])
        warnings = analysis_data.get("style_warnings", [])
        recency = analysis_data.get("recency", {})

        summary_text = html.escape(f"Parsed {len(citations)} citations, {len(refs)} reference entries, and {len(warnings)} style warnings.")
        story.append(Paragraph(f"<b>Summary:</b> {summary_text}", body_style))
        story.append(Spacer(1, 10))

        if warnings:
            story.append(Paragraph("Highest-Priority Findings", h2_style))
            for w in warnings[:15]:
                code = html.escape(str(w.get('code', 'STYLE')))
                msg_text = html.escape(str(w.get('message', '')))
                msg = f"• <b>[{code}]</b> {msg_text}"
                story.append(Paragraph(msg, body_style))

        retracted = [r for r in refs if r.get("status") == "retracted" or r.get("retraction_info", {}).get("is_retracted")]
        if retracted:
            story.append(Paragraph("Retracted Sources Flagged", h2_style))
            for r in retracted:
                raw_entry = html.escape(str(r.get('raw_entry', '')))
                ret_msg = f"• <font color='red'><b>RETRACTED:</b> {raw_entry}</font>"
                story.append(Paragraph(ret_msg, body_style))

        story.append(Paragraph("Publication Recency Breakdown", h2_style))
        p5 = html.escape(str(recency.get('within_5_years_percent', 0)))
        p10 = html.escape(str(recency.get('within_10_years_percent', 0)))
        avg_age = html.escape(str(recency.get('average_source_age_years', 'N/A')))
        rec_msg = f"Within 5 Years: {p5}% | Within 10 Years: {p10}% | Avg Age: {avg_age} yrs"
        story.append(Paragraph(rec_msg, body_style))

        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes

    except ImportError:
        logger.warning("reportlab package not installed. Generating valid minimal PDF fallback.")
        lines = [
            f"Summary: Parsed {len(analysis_data.get('citations', []))} citations, {len(analysis_data.get('references', []))} references, {len(analysis_data.get('style_warnings', []))} style warnings."
        ]
        for w in analysis_data.get("style_warnings", [])[:10]:
            lines.append(f"Style Warning [{w.get('code')}]: {w.get('message')}")
        return _generate_minimal_pdf_bytes("CitePilot Diagnostic Report", lines)


def _generate_minimal_pdf_bytes(title: str, lines: list) -> bytes:
    """Generates valid PDF 1.4 bytes without external PDF libraries (computes correct xref offsets)."""
    content_stream = f"BT /F1 14 Tf 40 750 Td 18 TL ({title}) Tj T*"
    for line in lines[:25]:
        clean = str(line).replace("(", "\\(").replace(")", "\\)")
        content_stream += f" ({clean}) Tj T*"
    content_stream += " ET"

    stream_len = len(content_stream)
    # Build PDF objects incrementally and track byte offsets for proper xref table
    objects = []
    # Object 1: Catalog
    obj1 = b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
    objects.append(obj1)
    # Object 2: Pages
    obj2 = b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
    objects.append(obj2)
    # Object 3: Page
    obj3 = b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
    objects.append(obj3)
    # Object 4: Content stream
    stream_bytes = content_stream.encode("latin1", errors="replace")
    obj4 = f"4 0 obj\n<< /Length {stream_len} >>\nstream\n".encode("latin1") + stream_bytes + b"\nendstream\nendobj\n"
    objects.append(obj4)
    # Object 5: Font
    obj5 = b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
    objects.append(obj5)

    header = b"%PDF-1.4\n"
    # Compute byte offsets: each object starts at cumulative offset after header
    offsets = [len(header)]
    for obj_bytes in objects:
        offsets.append(offsets[-1] + len(obj_bytes))

    xref_entries = b""
    for i in range(6):
        if i == 0:
            xref_entries += f"{offsets[0]:010d} 65535 f \n".encode("latin1")
        else:
            xref_entries += f"{offsets[i - 1]:010d} 00000 n \n".encode("latin1")

    trailer = b"trailer\n<< /Size 6 /Root 1 0 R >>\n"
    startxref_offset = len(header) + sum(len(o) for o in objects) + len(b"xref\n") + len(xref_entries) + len(trailer)

    pdf = (
        header +
        b"".join(objects) +
        b"xref\n" +
        xref_entries +
        trailer +
        f"startxref\n{startxref_offset}\n%%EOF".encode("latin1")
    )
    return pdf


def generate_redline_docx(text: str, analysis_data: Dict) -> bytes:
    """
    Generates a Redline DOCX manuscript with highlighted annotations and tracked style warnings.
    """
    doc = Document()
    doc.add_heading("CitePilot Redline Annotated Manuscript", level=1)

    warnings = analysis_data.get("style_warnings", [])
    citations = analysis_data.get("citations", [])
    claims = analysis_data.get("uncited_claims", [])

    doc.add_paragraph(
        f"Audit Summary: {len(warnings)} Style Warnings, {len(citations)} In-Text Citations, and {len(claims)} Uncited Claims Detected."
    )
    doc.add_paragraph("----------------------------------------------------------------------")

    # Map warnings by target_text for quick lookup
    warning_map = {}
    for w in warnings:
        target = w.get("target_text", "").strip()
        if target:
            warning_map[target] = w

    paragraphs = text.split("\n\n") if text else []
    for para_text in paragraphs:
        p = doc.add_paragraph()
        p_str = para_text.strip()
        if not p_str:
            continue

        # Check if paragraph contains flagged targets
        matched_warnings = [w for target, w in warning_map.items() if target in p_str]

        if matched_warnings:
            run = p.add_run(p_str)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            for mw in matched_warnings:
                comment_run = p.add_run(f"  [STYLE WARNING ({mw.get('code')}): {mw.get('message')}]")
                comment_run.bold = True
                comment_run.font.highlight_color = WD_COLOR_INDEX.PINK
        else:
            p.add_run(p_str)

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return docx_bytes
